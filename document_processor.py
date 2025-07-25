import streamlit as st
import PyPDF2
import docx
import io
from typing import List

class DocumentProcessor:
    """Handles document text extraction and chunking."""
    
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        """
        Initialize the document processor.
        
        Args:
            chunk_size: Maximum size of each text chunk
            chunk_overlap: Number of characters to overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def extract_text(self, uploaded_file) -> str:
        """
        Extract text from uploaded file based on its type.
        
        Args:
            uploaded_file: Streamlit uploaded file object
            
        Returns:
            Extracted text content
            
        Raises:
            Exception: If file type is not supported or extraction fails
        """
        file_type = uploaded_file.type
        
        try:
            if file_type == "application/pdf":
                return self._extract_from_pdf(uploaded_file)
            elif file_type == "text/plain":
                return self._extract_from_txt(uploaded_file)
            elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
                return self._extract_from_docx(uploaded_file)
            else:
                raise Exception(f"Unsupported file type: {file_type}")
                
        except Exception as e:
            raise Exception(f"Failed to extract text from {uploaded_file.name}: {str(e)}")
    
    def _extract_from_pdf(self, uploaded_file) -> str:
        """Extract text from PDF file."""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(uploaded_file.read()))
            text = ""
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                except Exception as e:
                    st.warning(f"Could not extract text from page {page_num + 1}: {str(e)}")
                    continue
            
            if not text.strip():
                raise Exception("No text content found in PDF")
                
            return text
            
        except Exception as e:
            raise Exception(f"PDF processing error: {str(e)}")
    
    def _extract_from_txt(self, uploaded_file) -> str:
        """Extract text from TXT file."""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    uploaded_file.seek(0)  # Reset file pointer
                    text = uploaded_file.read().decode(encoding)
                    return text
                except UnicodeDecodeError:
                    continue
            
            raise Exception("Could not decode text file with supported encodings")
            
        except Exception as e:
            raise Exception(f"TXT processing error: {str(e)}")
    
    def _extract_from_docx(self, uploaded_file) -> str:
        """Extract text from DOCX file."""
        try:
            doc = docx.Document(io.BytesIO(uploaded_file.read()))
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            text += cell.text + " "
                    text += "\n"
            
            if not text.strip():
                raise Exception("No text content found in DOCX")
                
            return text
            
        except Exception as e:
            raise Exception(f"DOCX processing error: {str(e)}")
    
    def chunk_text(self, text: str) -> List[str]:
        """
        Split text into overlapping chunks for better retrieval.
        
        Args:
            text: Input text to chunk
            
        Returns:
            List of text chunks
        """
        if not text.strip():
            return []
        
        # Clean and normalize text
        text = text.replace('\n\n', '\n').replace('\r', '').strip()
        
        chunks = []
        start = 0
        
        while start < len(text):
            # Find the end of the current chunk
            end = start + self.chunk_size
            
            # If we're not at the end of the text, try to end at a sentence boundary
            if end < len(text):
                # Look for sentence endings within the last 100 characters
                search_start = max(start + self.chunk_size - 100, start)
                sentence_ends = []
                
                for i in range(search_start, min(end + 50, len(text))):
                    if text[i] in '.!?':
                        # Check if it's likely end of sentence (followed by space or newline)
                        if i + 1 < len(text) and text[i + 1] in ' \n':
                            sentence_ends.append(i + 1)
                
                # Use the last sentence ending if found
                if sentence_ends:
                    end = sentence_ends[-1]
            
            # Extract chunk
            chunk = text[start:end].strip()
            
            if chunk:
                chunks.append(chunk)
            
            # Move start position (with overlap)
            start = end - self.chunk_overlap
            
            # Ensure we make progress
            if start <= 0:
                start = end
        
        # Remove very short chunks (less than 50 characters)
        chunks = [chunk for chunk in chunks if len(chunk) >= 50]
        
        return chunks
